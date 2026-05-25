from dataclasses import dataclass
from io import BytesIO

from app.core.config import settings
from app.core.ocr_capabilities import check_binary


class ResumeParseError(RuntimeError):
    pass


@dataclass(frozen=True)
class ParsedResume:
    text: str
    parser_version: str
    metadata: dict[str, str | int]


class ExtractionService:
    parser_version = "resume-extraction-0.3.0"

    def parse(self, payload: bytes, content_type: str) -> ParsedResume:
        if not payload:
            raise ResumeParseError("Document payload is empty")
        if len(payload) > settings.max_upload_bytes:
            raise ResumeParseError("Document payload exceeds configured upload limit")
        if content_type == "application/pdf":
            return self._parse_pdf(payload)
        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._parse_docx(payload)
        if content_type in {"image/png", "image/jpeg"}:
            return self._parse_image(payload)
        raise ResumeParseError(f"Unsupported content type: {content_type}")

    def _parse_pdf(self, payload: bytes) -> ParsedResume:
        text, page_count, method, metadata = self._extract_pdf_direct(payload)
        if len(text) >= settings.ocr_min_text_chars:
            return ParsedResume(
                text=text,
                parser_version=self.parser_version,
                metadata={"page_count": page_count, "method": method, **metadata},
            )

        if not settings.ocr_enabled:
            return ParsedResume(
                text=text,
                parser_version=self.parser_version,
                metadata={"page_count": page_count, "method": method, "ocr_skipped": "disabled", **metadata},
            )

        ocr_text, ocr_pages = self._ocr_pdf_pages(payload)
        combined = normalize_text("\n".join(part for part in [text, ocr_text] if part))
        return ParsedResume(
            text=combined,
            parser_version=self.parser_version,
            metadata={
                "page_count": page_count,
                "ocr_pages": ocr_pages,
                "method": "direct+ocr" if text else "ocr",
                **metadata,
            },
        )

    def _parse_docx(self, payload: bytes) -> ParsedResume:
        try:
            from docx import Document
        except Exception as exc:
            raise ResumeParseError(f"DOCX parser unavailable: {exc}") from exc

        document = Document(BytesIO(payload))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        text = normalize_text("\n".join(paragraphs))
        return ParsedResume(
            text=text,
            parser_version=self.parser_version,
            metadata={"paragraph_count": len(document.paragraphs)},
        )

    def _parse_image(self, payload: bytes) -> ParsedResume:
        if not settings.ocr_enabled:
            raise ResumeParseError("Image OCR is disabled")
        self._require_tesseract()
        try:
            import pytesseract
            from PIL import Image
        except Exception as exc:
            raise ResumeParseError(f"Image OCR dependencies unavailable: {exc}") from exc

        image = Image.open(BytesIO(payload))
        image.verify()
        image = Image.open(BytesIO(payload))
        text = normalize_text(
            pytesseract.image_to_string(image, timeout=settings.ocr_timeout_seconds)
        )
        return ParsedResume(
            text=text,
            parser_version=self.parser_version,
            metadata={"width": image.width, "height": image.height, "method": "ocr"},
        )

    def _extract_pdf_direct(self, payload: bytes) -> tuple[str, int, str, dict[str, str]]:
        try:
            import pdfplumber

            with pdfplumber.open(BytesIO(payload)) as pdf:
                page_count = len(pdf.pages)
                extract_pages = min(page_count, settings.ocr_max_pages)
                pages = [page.extract_text() or "" for page in pdf.pages[:extract_pages]]
                metadata = {
                    key: str(value)
                    for key, value in (pdf.metadata or {}).items()
                    if key.lower() in {"title", "author", "subject"} and value
                }
                if page_count > extract_pages:
                    metadata["truncated_pages"] = str(page_count - extract_pages)
                    metadata["extracted_pages"] = str(extract_pages)
            return normalize_text("\n".join(pages)), page_count, "pdfplumber", metadata
        except ResumeParseError:
            raise
        except Exception:
            pass

        try:
            import fitz
        except Exception as exc:
            raise ResumeParseError(f"PDF parser unavailable: {exc}") from exc

        try:
            document = fitz.open(stream=payload, filetype="pdf")
            try:
                page_count = document.page_count
                extract_pages = min(page_count, settings.ocr_max_pages)
                metadata = {
                    key: str(value)
                    for key, value in (document.metadata or {}).items()
                    if key.lower() in {"title", "author", "subject"} and value
                }
                if page_count > extract_pages:
                    metadata["truncated_pages"] = str(page_count - extract_pages)
                    metadata["extracted_pages"] = str(extract_pages)
                return normalize_text(
                    "\n".join(document[page_index].get_text() for page_index in range(extract_pages))
                ), page_count, "pymupdf", metadata
            finally:
                document.close()
        except ResumeParseError:
            raise
        except Exception as exc:
            raise ResumeParseError(f"PDF extraction failed: {exc}") from exc

    def _ocr_pdf_pages(self, payload: bytes) -> tuple[str, int]:
        self._require_tesseract()
        try:
            import fitz
            import pytesseract
            from PIL import Image
        except Exception as exc:
            raise ResumeParseError(f"PDF OCR dependencies unavailable: {exc}") from exc

        document = fitz.open(stream=payload, filetype="pdf")
        try:
            page_count = min(document.page_count, settings.ocr_max_pages)
            zoom = settings.ocr_pdf_render_dpi / 72
            matrix = fitz.Matrix(zoom, zoom)
            pages: list[str] = []
            for page_index in range(page_count):
                pixmap = document[page_index].get_pixmap(matrix=matrix, alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                pages.append(
                    pytesseract.image_to_string(image, timeout=settings.ocr_timeout_seconds)
                )
            return normalize_text("\n".join(pages)), page_count
        finally:
            document.close()

    def _require_tesseract(self) -> None:
        binary = check_binary("tesseract")
        if not binary.available:
            raise ResumeParseError("Tesseract OCR binary is unavailable; OCR fallback skipped")


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)
