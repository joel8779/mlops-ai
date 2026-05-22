from dataclasses import dataclass
from io import BytesIO

from docx import Document
from PIL import Image
from pypdf import PdfReader
import pytesseract


class ResumeParseError(RuntimeError):
    pass


@dataclass(frozen=True)
class ParsedResume:
    text: str
    parser_version: str
    metadata: dict[str, str | int]


class ResumeParser:
    parser_version = "resume-parser-0.1.0"

    def parse(self, payload: bytes, content_type: str) -> ParsedResume:
        if content_type == "application/pdf":
            return self._parse_pdf(payload)
        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._parse_docx(payload)
        if content_type in {"image/png", "image/jpeg"}:
            return self._parse_image(payload)
        raise ResumeParseError(f"Unsupported content type: {content_type}")

    def _parse_pdf(self, payload: bytes) -> ParsedResume:
        reader = PdfReader(BytesIO(payload))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = normalize_text("\n".join(pages))
        return ParsedResume(
            text=text,
            parser_version=self.parser_version,
            metadata={"page_count": len(reader.pages), "parser": "pypdf"},
        )

    def _parse_docx(self, payload: bytes) -> ParsedResume:
        document = Document(BytesIO(payload))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        text = normalize_text("\n".join(paragraphs))
        return ParsedResume(
            text=text,
            parser_version=self.parser_version,
            metadata={"paragraph_count": len(document.paragraphs), "parser": "python-docx"},
        )

    def _parse_image(self, payload: bytes) -> ParsedResume:
        image = Image.open(BytesIO(payload))
        text = normalize_text(pytesseract.image_to_string(image))
        return ParsedResume(
            text=text,
            parser_version=self.parser_version,
            metadata={"width": image.width, "height": image.height, "parser": "tesseract"},
        )


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    compact_lines = [line for line in lines if line]
    return "\n".join(compact_lines)
